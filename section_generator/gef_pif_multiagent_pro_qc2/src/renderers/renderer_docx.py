from __future__ import annotations
from typing import Dict, List, Any
from docx import Document
import os
import json

def _try_parse_table(body: str) -> Dict[str, Any] | None:
    try:
        obj = json.loads(body)
        if isinstance(obj, dict) and "table_data" in obj:
            return obj
    except Exception:
        pass
    return None

def _add_table(doc: Document, rows: List[List[str]], header: List[str] | None = None):
    if header:
        table = doc.add_table(rows=1, cols=len(header))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(header):
            hdr_cells[i].text = h
        for r in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(r):
                row_cells[i].text = val
    else:
        table = doc.add_table(rows=0, cols=len(rows[0]) if rows else 1)
        for r in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(r):
                row_cells[i].text = val

def _render_stakeholders_docx(doc: Document, obj: Dict[str, Any]):
    for group in obj.get("table_data", []):
        group_type = group.get("type", "Group")
        doc.add_heading(group_type, level=2)
        entries = group.get("entries", []) or []
        use_simple = any(isinstance(e, dict) and ("existing_activities" in e) for e in entries)
        header = ["Name", "Existing activities"] if use_simple else ["Name", "Activities", "Source URLs", "Confidence"]
        rows: List[List[str]] = []
        for entry in group.get("entries", []):
            name = entry.get("name", "")
            if use_simple:
                rows.append([name, entry.get("existing_activities", "")])
            else:
                activities = ", ".join(entry.get("activities", []) or [])
                urls = ", ".join(entry.get("source_urls", []) or [])
                conf = str(entry.get("confidence", ""))
                rows.append([name, activities, urls, conf])
        _add_table(doc, rows, header=header)

def _render_simple_table_docx(doc: Document, obj: Dict[str, Any], columns: List[str]):
    header = columns
    rows: List[List[str]] = []
    for row in obj.get("table_data", []):
        rows.append([str(row.get(col, "")) for col in columns])
    _add_table(doc, rows, header=header)
    if "summary" in obj and isinstance(obj["summary"], str) and obj["summary"].strip():
        doc.add_paragraph(obj["summary"].strip())

def _render_other_baseline_initiatives_docx(doc: Document, obj: Dict[str, Any]):
    aliases: Dict[str, List[str]] = {
        "program": ["program_project", "program/project", "program", "project"],
        "entities": ["leading_entities", "Leading Ministry and Supporting Entities"],
        "desc": ["description"],
        "duration": ["duration"],
        "value": ["value_usd", "value"],
        "relation": ["relation_to_etf", "Relationship with ETF and the transparency system"],
    }
    def pick(row: Dict[str, Any], keys: List[str]) -> str:
        for k in keys:
            v = row.get(k)
            if isinstance(v, (str, int, float)) and str(v).strip():
                return str(v)
        return ""
    header = ["Program/Project", "Leading Entities", "Description", "Duration", "Value (USD)", "Relation to ETF"]
    rows: List[List[str]] = []
    for row in obj.get("table_data", []):
        rows.append([
            pick(row, aliases["program"]),
            pick(row, aliases["entities"]),
            pick(row, aliases["desc"]),
            pick(row, aliases["duration"]),
            pick(row, aliases["value"]),
            pick(row, aliases["relation"]),
        ])
    _add_table(doc, rows, header=header)

def assemble_document_docx(parts: Dict[str, str], titles: Dict[str, str], country: str, out_path: str) -> str:
    os.makedirs(os.path.dirname(out_path) or '.', exist_ok=True)
    order = [
        "rationale_intro","paris_etf","climate_transparency_country",
        "baseline_national_tf_header","baseline_institutional","baseline_policy",
        "baseline_stakeholders","baseline_unfccc_reporting",
        "module_header","module_ghg","module_adaptation","module_ndc_tracking","module_support",
        "other_baseline_initiatives","key_barriers","barrier1","barrier2","barrier3",
        "appendix_quality"
    ]
    doc = Document()
    doc.add_heading(f"GEF-8 PROJECT IDENTIFICATION FORM (PIF) — {country}", 0)
    for key in order:
        body = (parts.get(key) or "").strip()
        if not body: continue
        doc.add_heading(titles.get(key, key), level=1)
        table_obj = _try_parse_table(body)
        if key == "baseline_stakeholders" and table_obj:
            _render_stakeholders_docx(doc, table_obj)
            continue
        if key == "baseline_unfccc_reporting" and table_obj:
            _render_simple_table_docx(doc, table_obj, ["year", "report", "comment"])
            continue
        if key == "other_baseline_initiatives" and table_obj:
            _render_other_baseline_initiatives_docx(doc, table_obj)
            continue
        for p in [p.strip() for p in body.split("\n") if p.strip()]:
            doc.add_paragraph(p)
    doc.save(out_path)
    return out_path
