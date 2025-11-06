from __future__ import annotations
from typing import Dict, List
from docx import Document
import os

def assemble_document_docx(parts: Dict[str, str], titles: Dict[str, str], country: str, out_path: str) -> str:
    os.makedirs(os.path.dirname(out_path) or '.', exist_ok=True)
    order = [
        "rationale_intro","paris_etf","climate_transparency_country",
        "baseline_national_tf_header","baseline_institutional","baseline_policy",
        "baseline_stakeholders","baseline_unfccc_reporting",
        "module_header","module_ghg","module_adaptation","module_ndc_tracking","module_support",
        "other_baseline_initiatives","key_barriers",
        "appendix_quality"
    ]
    doc = Document()
    doc.add_heading(f"GEF-8 PROJECT IDENTIFICATION FORM (PIF) — {country}", 0)
    for key in order:
        body = (parts.get(key) or "").strip()
        if not body: continue
        doc.add_heading(titles.get(key, key), level=1)
        for p in [p.strip() for p in body.split("\n") if p.strip()]:
            doc.add_paragraph(p)
    doc.save(out_path)
    return out_path
